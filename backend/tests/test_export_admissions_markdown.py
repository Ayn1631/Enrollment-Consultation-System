from __future__ import annotations

from pathlib import Path

from docx import Document as DocxDocument

from app.admissions_kb.parsers import (
    export_admission_markdown_documents,
    parse_admission_docx,
    render_admission_markdown,
)


def test_parse_admission_docx_extracts_metadata_and_blocks(tmp_path: Path):
    doc_path = tmp_path / "college.docx"
    document = DocxDocument()
    document.add_paragraph("数学与信息科学学院")
    document.add_paragraph("（招生咨询电话：0371-62506034，62506038）")
    document.add_paragraph("（学院网址：https://sxxy.zut.edu.cn/）")
    document.add_paragraph("学院简介")
    document.add_paragraph("学院现有教职工76人。")
    document.save(doc_path)

    parsed = parse_admission_docx(doc_path)

    assert parsed.title == "数学与信息科学学院"
    assert parsed.metadata["contact_phone"].startswith("0371-62506034")
    assert parsed.metadata["website"] == "https://sxxy.zut.edu.cn/"
    assert "学院简介" in parsed.content_blocks


def test_render_admission_markdown_keeps_zyit_header(tmp_path: Path):
    doc_path = tmp_path / "policy.docx"
    document = DocxDocument()
    document.add_paragraph("中原工学院2025年普通本科招生章程")
    document.add_paragraph("第一章 总则")
    document.add_paragraph("为了保证招生工作顺利进行，制定本章程。")
    document.save(doc_path)

    parsed = parse_admission_docx(doc_path)
    markdown = render_admission_markdown(parsed)

    assert markdown.startswith(f"# 原文（来源：{doc_path}）")
    assert "网页标题：中原工学院2025年普通本科招生章程" in markdown
    assert "抓取时间：" in markdown


def test_export_admission_markdown_documents_writes_markdown_and_json(tmp_path: Path):
    source_root = tmp_path / "source"
    output_dir = tmp_path / "zyit"
    parsed_dir = tmp_path / "parsed"
    source_root.mkdir(parents=True)
    doc_path = source_root / "sample.docx"
    document = DocxDocument()
    document.add_paragraph("自动化与电气工程学院")
    document.add_paragraph("学院简介")
    document.add_paragraph("自动化与电气工程学院前身是电气系。")
    document.save(doc_path)

    exported = export_admission_markdown_documents(
        source_root=source_root,
        output_dir=output_dir,
        parsed_dir=parsed_dir,
    )

    assert len(exported) == 1
    markdown = exported[0].read_text(encoding="utf-8")
    json_path = parsed_dir / f"{exported[0].stem}.json"
    assert "自动化与电气工程学院" in markdown
    assert json_path.exists()


def test_export_admission_markdown_real_docx(runtime_settings, tmp_path: Path):
    source_root = Path(runtime_settings.docs_dir).parent / "招生资料" / "25"
    target = source_root / "学院信息" / "8、数学与信息科学学院招生宣传材料.docx"
    if not target.exists():
        raise AssertionError(f"真实学院文档不存在: {target}")

    exported = export_admission_markdown_documents(
        source_root=target.parent,
        output_dir=tmp_path / "zyit",
        parsed_dir=tmp_path / "parsed",
    )

    assert exported
    target_export = next((path for path in exported if "数学与信息科学学院" in path.read_text(encoding="utf-8")), None)
    assert target_export is not None
    text = target_export.read_text(encoding="utf-8")
    assert "数学与信息科学学院" in text
    assert "学院简介" in text
