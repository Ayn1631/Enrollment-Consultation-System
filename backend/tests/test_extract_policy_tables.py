from __future__ import annotations

from pathlib import Path

from docx import Document as DocxDocument
from openpyxl import load_workbook

from app.admissions_kb.parsers import MAJOR_CATALOG_FIELDS, extract_policy_tables_from_docx, export_policy_tables_to_excel


def test_extract_policy_tables_handles_duplicate_headers(tmp_path: Path):
    doc_path = tmp_path / "policy.docx"
    document = DocxDocument()
    table = document.add_table(rows=4, cols=10)
    title_row = table.rows[0].cells
    for cell in title_row:
        cell.text = "中原工学院2025年专业情况汇总表"
    header_cells = table.rows[1].cells
    headers = ["序号", "专业代码", "专业代码", "专业名称", "学制", "学费（元）", "选考科目", "学位授予门类", "所在院系", "所在院系"]
    for cell, header in zip(header_cells, headers):
        cell.text = header
    first_data = ["1", "0801", "0801", "测试专业", "四年", "5000", "物理+化学", "工学", "测试学院", "测试学院"]
    second_data = ["2", "0802", "0802", "第二专业", "四年", "5500", "不限", "理学", "第二学院", "第二学院"]
    for row_cells, row_values in zip(table.rows[2:], [first_data, second_data]):
        for cell, value in zip(row_cells.cells, row_values):
            cell.text = value
    document.save(doc_path)

    tables = extract_policy_tables_from_docx(doc_path)

    assert len(tables) == 1
    assert tables[0].dataset == "major_catalog"
    assert tables[0].rows[0]["专业名称"] == "测试专业"
    assert tables[0].rows[0]["所在院系"] == "测试学院"
    assert all(field in tables[0].rows[0] for field in MAJOR_CATALOG_FIELDS)


def test_export_policy_tables_to_excel_writes_expected_headers(tmp_path: Path):
    doc_path = tmp_path / "policy.docx"
    out_path = tmp_path / "derived.xlsx"
    document = DocxDocument()
    table = document.add_table(rows=3, cols=8)
    for cell in table.rows[0].cells:
        cell.text = "测试政策表"
    headers = ["序号", "专业代码", "专业名称", "学制", "学费（元）", "选考科目", "学位授予门类", "所在院系"]
    for cell, header in zip(table.rows[1].cells, headers):
        cell.text = header
    values = ["1", "0801", "测试专业", "四年", "5000", "物理+化学", "工学", "测试学院"]
    for cell, value in zip(table.rows[2].cells, values):
        cell.text = value
    document.save(doc_path)

    export_policy_tables_to_excel(input_path=doc_path, output_path=out_path)

    workbook = load_workbook(out_path, read_only=True, data_only=True)
    sheet = workbook["major_catalog"]
    rows = list(sheet.iter_rows(values_only=True))
    assert rows[0][:8] == tuple(MAJOR_CATALOG_FIELDS)
    assert rows[1][2] == "测试专业"
    assert rows[1][8] == doc_path.name


def test_extract_policy_tables_from_real_admission_doc(runtime_settings):
    doc_path = Path(runtime_settings.docs_dir).parent / "招生资料" / "25" / "中原工学院2025年普通本科招生章程.docx"
    if not doc_path.exists():
        raise AssertionError(f"真实章程文件不存在: {doc_path}")

    tables = extract_policy_tables_from_docx(doc_path)

    assert tables
    major_catalog = next((table for table in tables if table.dataset == "major_catalog"), None)
    assert major_catalog is not None
    assert "专业情况汇总表" in major_catalog.title
    assert len(major_catalog.rows) >= 60
    assert {"专业代码", "专业名称", "学费（元）", "选考科目", "所在院系"}.issubset(major_catalog.rows[0].keys())
