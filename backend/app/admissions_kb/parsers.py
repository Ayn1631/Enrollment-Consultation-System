from __future__ import annotations

from dataclasses import dataclass
from datetime import datetime
import json
from pathlib import Path
import re

from docx import Document as DocxDocument
from openpyxl import Workbook
from openpyxl import load_workbook
from PyPDF2 import PdfReader


MAJOR_CATALOG_FIELDS = [
    "序号",
    "专业代码",
    "专业名称",
    "学制",
    "学费（元）",
    "选考科目",
    "学位授予门类",
    "所在院系",
]
TRACE_FIELDS = [
    "source_file",
    "source_doc",
    "source_table_title",
    "source_row_no",
    "extract_time",
]
HEADER_ALIASES = {
    "序号": "序号",
    "专业代码": "专业代码",
    "专业名称": "专业名称",
    "学制": "学制",
    "学费": "学费（元）",
    "学费（元）": "学费（元）",
    "选考科目": "选考科目",
    "学位授予门类": "学位授予门类",
    "学位": "学位授予门类",
    "所在院系": "所在院系",
    "院系": "所在院系",
}


@dataclass(slots=True)
class ExtractedPolicyTable:
    dataset: str
    title: str
    source_file: str
    source_doc: str
    rows: list[dict[str, str]]


@dataclass(slots=True)
class ParsedAdmissionDocument:
    source_file: str
    source_path: str
    source_type: str
    title: str
    publish_date: str
    grab_date: str
    source_url: str
    content_blocks: list[str]
    warnings: list[str]
    metadata: dict[str, str]

    def to_json(self) -> str:
        payload = {
            "source_file": self.source_file,
            "source_path": self.source_path,
            "source_type": self.source_type,
            "title": self.title,
            "publish_date": self.publish_date,
            "grab_date": self.grab_date,
            "source_url": self.source_url,
            "content_blocks": self.content_blocks,
            "warnings": self.warnings,
            "metadata": self.metadata,
        }
        return json.dumps(payload, ensure_ascii=False, indent=2)


def extract_policy_tables_from_docx(input_path: Path) -> list[ExtractedPolicyTable]:
    document = DocxDocument(str(input_path))
    tables: list[ExtractedPolicyTable] = []
    for index, table in enumerate(document.tables, start=1):
        extracted = _extract_single_table(table_rows=_table_to_rows(table), input_path=input_path, table_index=index)
        if extracted is not None:
            tables.append(extracted)
    return tables


def export_policy_tables_to_excel(input_path: Path, output_path: Path) -> list[ExtractedPolicyTable]:
    extracted_tables = extract_policy_tables_from_docx(input_path)
    if not extracted_tables:
        raise ValueError(f"未在 {input_path} 中识别到可结构化的政策表格")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    workbook = Workbook()
    first_sheet = True
    for table in extracted_tables:
        sheet = workbook.active if first_sheet else workbook.create_sheet()
        first_sheet = False
        sheet.title = _safe_sheet_title(table.dataset)
        headers = MAJOR_CATALOG_FIELDS + TRACE_FIELDS
        sheet.append(headers)
        for row in table.rows:
            sheet.append([row.get(field, "") for field in headers])
    workbook.save(output_path)
    return extracted_tables


def parse_admission_docx(input_path: Path) -> ParsedAdmissionDocument:
    document = DocxDocument(str(input_path))
    paragraphs = [_normalize_cell(paragraph.text) for paragraph in document.paragraphs]
    content_blocks = [block for block in _merge_docx_blocks(paragraphs) if block]
    title = _resolve_docx_title(content_blocks, input_path)
    publish_date = _extract_publish_date_from_text("\n".join(content_blocks[:40]))
    warnings: list[str] = []
    if input_path.name == "中原工学院2025年普通本科招生章程.docx":
        content_blocks = _inject_policy_table_note(content_blocks)
    if not content_blocks:
        warnings.append("未抽取到有效正文段落")
    metadata = _extract_common_metadata(content_blocks, input_path)
    return ParsedAdmissionDocument(
        source_file=input_path.name,
        source_path=str(input_path),
        source_type="docx",
        title=title,
        publish_date=publish_date,
        grab_date=datetime.now().date().isoformat(),
        source_url=str(input_path),
        content_blocks=content_blocks,
        warnings=warnings,
        metadata=metadata,
    )


def parse_admission_pdf(input_path: Path) -> ParsedAdmissionDocument:
    reader = PdfReader(str(input_path))
    pages: list[str] = []
    warnings: list[str] = []
    for index, page in enumerate(reader.pages, start=1):
        text = _normalize_pdf_text(page.extract_text() or "")
        if not text:
            continue
        pages.append(f"第{index}页\n{text}")
    joined = "\n".join(pages)
    if not pages:
        warnings.append("PDF 未抽取到可用文本")
    elif _looks_garbled(joined):
        warnings.append("PDF 文本疑似存在乱码，请考虑后续补 OCR")
    metadata = _extract_common_metadata(pages, input_path)
    return ParsedAdmissionDocument(
        source_file=input_path.name,
        source_path=str(input_path),
        source_type="pdf",
        title=input_path.stem,
        publish_date=_extract_publish_date_from_text(joined[:2000]),
        grab_date=datetime.now().date().isoformat(),
        source_url=str(input_path),
        content_blocks=pages,
        warnings=warnings,
        metadata=metadata,
    )


def render_admission_markdown(document: ParsedAdmissionDocument) -> str:
    lines = [
        f"# 原文（来源：{document.source_url}）",
        f"网页标题：{document.title}",
    ]
    if document.publish_date:
        lines.append(f"发布时间：{document.publish_date}")
    lines.append(f"抓取时间：{document.grab_date}")
    lines.append("")
    if document.metadata.get("contact_phone"):
        lines.append(f"联系方式：{document.metadata['contact_phone']}")
        lines.append("")
    if document.metadata.get("website"):
        lines.append(f"网址：{document.metadata['website']}")
        lines.append("")
    if document.warnings:
        lines.append(f"解析提示：{'；'.join(document.warnings)}")
        lines.append("")
    for block in document.content_blocks:
        if _is_heading_block(block):
            lines.append(f"## {block}")
        else:
            lines.append(block)
        lines.append("")
    return "\n".join(lines).strip() + "\n"


def export_admission_markdown_documents(
    *,
    source_root: Path,
    output_dir: Path,
    parsed_dir: Path,
) -> list[Path]:
    output_dir.mkdir(parents=True, exist_ok=True)
    parsed_dir.mkdir(parents=True, exist_ok=True)
    _cleanup_previous_exports(output_dir=output_dir, parsed_dir=parsed_dir)
    exported: list[Path] = []
    docx_paths = [path for path in sorted(source_root.rglob("*.docx")) if path.is_file()]
    pdf_paths = [path for path in sorted(source_root.rglob("*.pdf")) if path.is_file()]

    raw_documents: list[ParsedAdmissionDocument] = []
    for path in docx_paths:
        raw_documents.append(parse_admission_docx(path))
    for path in pdf_paths:
        raw_documents.append(parse_admission_pdf(path))

    for index, document in enumerate(raw_documents, start=19):
        stem = f"{index:02d}-{_build_markdown_name(document)}"
        markdown_path = output_dir / f"{stem}.md"
        parsed_path = parsed_dir / f"{stem}.json"
        markdown_path.write_text(render_admission_markdown(document), encoding="utf-8")
        parsed_path.write_text(document.to_json(), encoding="utf-8")
        exported.append(markdown_path)
    return exported


def load_major_catalog_rows(input_path: Path, *, source_dataset: str) -> list[dict[str, str]]:
    workbook = load_workbook(str(input_path), read_only=True, data_only=True)
    preferred_sheet = next((sheet for sheet in workbook.worksheets if _sheet_has_major_columns(sheet)), workbook.worksheets[0])
    rows = list(preferred_sheet.iter_rows(values_only=True))
    header_index = _find_first_non_empty_row(rows)
    if header_index is None:
        return []
    header_map = _build_header_map([str(value).strip() if value is not None else "" for value in rows[header_index]])
    results: list[dict[str, str]] = []
    for row_no, values in enumerate(rows[header_index + 1 :], start=header_index + 2):
        raw_row = [str(value).strip() if value is not None else "" for value in values]
        record = _build_major_catalog_row(
            raw_row=raw_row,
            header_map=header_map,
            source_file=input_path.name,
            source_doc=input_path.stem,
            source_table_title=preferred_sheet.title,
            source_row_no=str(row_no),
            extract_time=datetime.now().isoformat(timespec="seconds"),
        )
        if record is None:
            continue
        record["academic_year"] = _extract_primary_year(f"{input_path.stem} {' '.join(raw_row)}")
        record["source_dataset"] = source_dataset
        record["major_code"] = record.pop("专业代码", "")
        record["major_name"] = record.pop("专业名称", "")
        record["duration"] = record.pop("学制", "")
        record["tuition"] = record.pop("学费（元）", "")
        record["exam_subjects"] = record.pop("选考科目", "")
        record["degree_type"] = record.pop("学位授予门类", "")
        record["college_name"] = record.pop("所在院系", "")
        record["evidence_text"] = "；".join(
            f"{label}：{record.get(field, '')}"
            for label, field in (
                ("专业代码", "major_code"),
                ("专业名称", "major_name"),
                ("学制", "duration"),
                ("学费（元）", "tuition"),
                ("选考科目", "exam_subjects"),
                ("学位授予门类", "degree_type"),
                ("所在院系", "college_name"),
            )
            if record.get(field)
        )
        results.append(record)
    return results


def load_faq_seed_rows(input_path: Path) -> list[dict[str, str]]:
    workbook = load_workbook(str(input_path), read_only=True, data_only=True)
    sheet = workbook.worksheets[0]
    rows = list(sheet.iter_rows(values_only=True))
    header_index = _find_row_index_containing(rows, "问题")
    if header_index is None:
        return []
    headers = [str(value).strip() if value is not None else "" for value in rows[header_index]]
    column_map = {
        "tag_name": _find_column_index(headers, {"标签"}),
        "question_no": _find_column_index(headers, {"序号"}),
        "question": _find_column_index(headers, {"问题"}),
        "answer": _find_column_index(headers, {"答案（建议全面完善答案，考生提问后可以让其更全面的了解本问题，让考生咨询过程体验更佳。）", "答案"}),
    }
    results: list[dict[str, str]] = []
    for values in rows[header_index + 1 :]:
        row = [str(value).strip() if value is not None else "" for value in values]
        question = _value_by_index(row, column_map["question"])
        answer = _value_by_index(row, column_map["answer"])
        if not question or not answer:
            continue
        results.append(
            {
                "source_file": input_path.name,
                "tag_name": _value_by_index(row, column_map["tag_name"]),
                "question_no": _value_by_index(row, column_map["question_no"]),
                "question": question,
                "answer": answer,
                "retrieval_priority": "low",
            }
        )
    return results


def load_score_line_rows(input_path: Path) -> list[dict[str, str]]:
    try:
        import xlrd
    except Exception as exc:  # noqa: BLE001
        raise RuntimeError("缺少 xlrd，无法解析 xls 录取分数线文件") from exc

    workbook = xlrd.open_workbook(str(input_path))
    sheet = workbook.sheet_by_index(0)
    rows = [[_normalize_xls_value(sheet.cell_value(row_idx, col_idx)) for col_idx in range(sheet.ncols)] for row_idx in range(sheet.nrows)]
    header_index = _find_score_header_index(rows)
    if header_index is None:
        return []
    headers = rows[header_index]
    aliases = {
        "年份": {"年份", "年度"},
        "省份": {"省份", "生源省份", "省市"},
        "批次": {"批次", "录取批次"},
        "category": {"科类", "选科", "科类/选科"},
        "major_name": {"专业", "专业名称"},
        "min_score": {"最低分", "投档最低分", "录取最低分"},
        "min_rank": {"最低位次", "投档最低位次", "录取最低位次", "位次"},
    }
    column_map = {key: _find_column_index(headers, labels) for key, labels in aliases.items()}
    results: list[dict[str, str]] = []
    for row_no, values in enumerate(rows[header_index + 1 :], start=header_index + 2):
        if not any(values):
            continue
        major_name = _value_by_index(values, column_map["major_name"])
        if not major_name:
            continue
        evidence = "；".join(f"{headers[idx]}：{value}" for idx, value in enumerate(values) if value)
        results.append(
            {
                "source_dataset": "score_lines",
                "source_file": input_path.name,
                "source_sheet": sheet.name,
                "source_row_no": str(row_no),
                "year": _value_by_index(values, column_map["年份"]) or _extract_primary_year(f"{input_path.stem} {evidence}"),
                "province": _value_by_index(values, column_map["省份"]),
                "batch": _value_by_index(values, column_map["批次"]),
                "category": _value_by_index(values, column_map["category"]),
                "major_name": major_name,
                "min_score": _value_by_index(values, column_map["min_score"]),
                "min_rank": _value_by_index(values, column_map["min_rank"]),
                "evidence_text": evidence,
            }
        )
    return results


def _extract_single_table(
    *,
    table_rows: list[list[str]],
    input_path: Path,
    table_index: int,
) -> ExtractedPolicyTable | None:
    if not table_rows:
        return None
    title = _extract_table_title(table_rows)
    header_index = _find_header_index(table_rows)
    if header_index is None:
        return None
    header_map = _build_header_map(table_rows[header_index])
    if not _is_major_catalog_table(header_map):
        return None

    extracted_rows: list[dict[str, str]] = []
    extract_time = datetime.now().isoformat(timespec="seconds")
    for row_idx, row_values in enumerate(table_rows[header_index + 1 :], start=header_index + 2):
        row = _build_major_catalog_row(
            raw_row=row_values,
            header_map=header_map,
            source_file=input_path.name,
            source_doc=input_path.stem,
            source_table_title=title or f"table_{table_index}",
            source_row_no=str(row_idx),
            extract_time=extract_time,
        )
        if row is None:
            continue
        extracted_rows.append(row)
    if not extracted_rows:
        return None
    return ExtractedPolicyTable(
        dataset="major_catalog",
        title=title or f"table_{table_index}",
        source_file=input_path.name,
        source_doc=input_path.stem,
        rows=extracted_rows,
    )


def _table_to_rows(table) -> list[list[str]]:
    rows: list[list[str]] = []
    for row in table.rows:
        values = [_normalize_cell(cell.text) for cell in row.cells]
        if any(values):
            rows.append(values)
    return rows


def _normalize_cell(value: str) -> str:
    return re.sub(r"\s+", " ", (value or "").strip())


def _extract_table_title(table_rows: list[list[str]]) -> str:
    if not table_rows:
        return ""
    first_row = [item for item in table_rows[0] if item]
    if not first_row:
        return ""
    deduped = list(dict.fromkeys(first_row))
    return deduped[0] if len(deduped) == 1 else " ".join(deduped)


def _find_header_index(table_rows: list[list[str]]) -> int | None:
    best_idx: int | None = None
    best_score = -1
    for idx, row in enumerate(table_rows[:5]):
        normalized = [_normalize_header(cell) for cell in row]
        score = sum(1 for item in normalized if item in MAJOR_CATALOG_FIELDS)
        if score > best_score:
            best_idx = idx
            best_score = score
    return best_idx if best_score >= 4 else None


def _normalize_header(value: str) -> str:
    text = _normalize_cell(value)
    return HEADER_ALIASES.get(text, text)


def _build_header_map(header_row: list[str]) -> dict[str, int]:
    mapping: dict[str, int] = {}
    for idx, value in enumerate(header_row):
        normalized = _normalize_header(value)
        if normalized in MAJOR_CATALOG_FIELDS and normalized not in mapping:
            mapping[normalized] = idx
    return mapping


def _is_major_catalog_table(header_map: dict[str, int]) -> bool:
    required = {"专业代码", "专业名称", "学制", "学费（元）", "选考科目", "学位授予门类", "所在院系"}
    return required.issubset(set(header_map))


def _build_major_catalog_row(
    *,
    raw_row: list[str],
    header_map: dict[str, int],
    source_file: str,
    source_doc: str,
    source_table_title: str,
    source_row_no: str,
    extract_time: str,
) -> dict[str, str] | None:
    row: dict[str, str] = {}
    for field in MAJOR_CATALOG_FIELDS:
        value = raw_row[header_map[field]] if header_map[field] < len(raw_row) else ""
        row[field] = _normalize_cell(value)
    if not row.get("专业名称"):
        return None
    if row.get("专业名称") == "专业名称":
        return None
    row.update(
        {
            "source_file": source_file,
            "source_doc": source_doc,
            "source_table_title": source_table_title,
            "source_row_no": source_row_no,
            "extract_time": extract_time,
        }
    )
    return row


def _safe_sheet_title(value: str) -> str:
    cleaned = re.sub(r"[\[\]\*:/\\\?]", "_", value)
    return cleaned[:31] or "sheet1"


def _merge_docx_blocks(paragraphs: list[str]) -> list[str]:
    blocks: list[str] = []
    for paragraph in paragraphs:
        if not paragraph:
            continue
        if paragraph in {"学院简介", "专业介绍", "核心课程", "师资力量", "培养目标", "考研就业", "获奖证书"}:
            blocks.append(paragraph)
            continue
        if re.match(r"^[一二三四五六七八九十]+、", paragraph) or re.match(r"^\d+[、.]", paragraph):
            blocks.append(paragraph)
            continue
        if len(paragraph) <= 24 and any(token in paragraph for token in ("学院", "专业", "简介", "规则", "计划", "总则", "录取")):
            blocks.append(paragraph)
            continue
        if blocks and not _is_heading_block(blocks[-1]) and len(blocks[-1]) < 900:
            blocks[-1] = f"{blocks[-1]}\n{paragraph}"
            continue
        blocks.append(paragraph)
    return blocks


def _resolve_docx_title(content_blocks: list[str], input_path: Path) -> str:
    if content_blocks:
        candidate = content_blocks[0].splitlines()[0].strip()
        if candidate:
            return candidate
    return input_path.stem


def _extract_publish_date_from_text(text: str) -> str:
    patterns = [
        r"(20\d{2})年(\d{1,2})月(\d{1,2})日",
        r"(20\d{2})[-/](\d{1,2})[-/](\d{1,2})",
        r"(20\d{2})年(\d{1,2})月",
    ]
    for pattern in patterns:
        matched = re.search(pattern, text)
        if not matched:
            continue
        parts = [segment.zfill(2) for segment in matched.groups()]
        if len(parts) == 3:
            return f"{parts[0]}-{parts[1]}-{parts[2]}"
        if len(parts) == 2:
            return f"{parts[0]}-{parts[1]}-01"
    return ""


def _inject_policy_table_note(content_blocks: list[str]) -> list[str]:
    note = "说明：文中“中原工学院2025年专业情况汇总表”已抽离为结构化附表文件，可由结构化工具查询。"
    if any("专业情况汇总表" in block for block in content_blocks):
        return [note] + [block for block in content_blocks if "专业情况汇总表" not in block]
    return [note] + content_blocks


def _extract_common_metadata(content_blocks: list[str], input_path: Path) -> dict[str, str]:
    text = "\n".join(content_blocks[:60])
    phone = _search_first(text, r"((?:0\d{2,3}-)?\d{7,8}(?:、(?:0\d{2,3}-)?\d{7,8})*)")
    website = _search_first(text, r"(https?://[^\s）)]+)")
    return {
        "source_name": input_path.stem,
        "contact_phone": phone,
        "website": website,
    }


def _search_first(text: str, pattern: str) -> str:
    matched = re.search(pattern, text)
    return matched.group(1).strip() if matched else ""


def _normalize_pdf_text(value: str) -> str:
    return re.sub(r"\n{3,}", "\n\n", (value or "").replace("\x00", "").strip())


def _looks_garbled(text: str) -> bool:
    if not text:
        return False
    weird = sum(1 for char in text if char in "�□◌")
    return weird / max(len(text), 1) > 0.01 or "unknown widths" in text.lower()


def _is_heading_block(text: str) -> bool:
    stripped = text.strip()
    if "\n" in stripped:
        return False
    return bool(
        stripped
        and (
            stripped in {"学院简介", "专业介绍", "核心课程", "师资力量", "培养目标", "考研就业", "获奖证书"}
            or re.match(r"^[一二三四五六七八九十]+、", stripped)
            or re.match(r"^\d+[、.]", stripped)
            or len(stripped) <= 24 and any(token in stripped for token in ("学院", "专业", "总则", "规则", "计划", "章", "条"))
        )
    )


def _build_markdown_name(document: ParsedAdmissionDocument) -> str:
    if document.source_type == "pdf":
        return f"{_sanitize_filename(document.title)}（原始资料导出）"
    title = document.title.strip() or document.metadata.get("source_name", "")
    if "招生章程" in title or "报考指南" in title:
        return f"{_sanitize_filename(title)}（原始资料导出）"
    if "学院" in title and "招生宣传" not in title:
        return f"学院介绍-{_sanitize_filename(title)}"
    return f"{_sanitize_filename(title)}（原始资料导出）"


def _sanitize_filename(value: str) -> str:
    cleaned = re.sub(r'[<>:"/\\|?*]', "_", value.strip())
    cleaned = re.sub(r"\s+", "", cleaned)
    return cleaned[:80] or "未命名文档"


def _cleanup_previous_exports(*, output_dir: Path, parsed_dir: Path) -> None:
    markdown_pattern = re.compile(r"^\d{2}-(学院介绍-|.*原始资料导出).+\.md$")
    json_pattern = re.compile(r"^\d{2}-(学院介绍-|.*原始资料导出).+\.json$")
    for path in output_dir.glob("*.md"):
        if markdown_pattern.match(path.name):
            path.unlink(missing_ok=True)
    for path in parsed_dir.glob("*.json"):
        if json_pattern.match(path.name):
            path.unlink(missing_ok=True)


def _sheet_has_major_columns(sheet) -> bool:
    rows = list(sheet.iter_rows(values_only=True, max_row=3))
    for row in rows:
        headers = [str(value).strip() if value is not None else "" for value in row]
        header_map = _build_header_map(headers)
        if _is_major_catalog_table(header_map):
            return True
    return False


def _find_first_non_empty_row(rows: list[list[object]]) -> int | None:
    for idx, row in enumerate(rows):
        if any((str(value).strip() if value is not None else "") for value in row):
            return idx
    return None


def _find_row_index_containing(rows: list[list[object]], keyword: str) -> int | None:
    for idx, row in enumerate(rows):
        values = [str(value).strip() if value is not None else "" for value in row]
        if keyword in values:
            return idx
    return None


def _find_column_index(headers: list[str], labels: set[str]) -> int:
    for idx, header in enumerate(headers):
        if header in labels:
            return idx
    return -1


def _value_by_index(row: list[str], index: int) -> str:
    if index < 0 or index >= len(row):
        return ""
    return row[index]


def _extract_primary_year(text: str) -> str:
    matched = re.search(r"(20\d{2})", text)
    return matched.group(1) if matched else ""


def _normalize_xls_value(value: object) -> str:
    text = str(value).strip()
    if text.endswith(".0"):
        return text[:-2]
    return text


def _find_score_header_index(rows: list[list[str]]) -> int | None:
    expected_tokens = {"年份", "省份", "专业", "最低分", "位次", "批次", "录取批次"}
    best_idx: int | None = None
    best_score = -1
    for idx, row in enumerate(rows[:8]):
        score = sum(1 for cell in row if any(token in cell for token in expected_tokens))
        if score > best_score:
            best_idx = idx
            best_score = score
    return best_idx if best_score >= 2 else None
